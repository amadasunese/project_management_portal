from flask import Blueprint, render_template, request, redirect, url_for, abort, flash, session
from flask_weasyprint import HTML, render_pdf
from models import db
from models import Project, Activity, Report, FinancialReport
from forms import project_choices, ProjectForm, ReportForm, ActivityForm
from datetime import datetime
from werkzeug.utils import secure_filename
import os
from flask import Response
import pandas as pd
from io import BytesIO
from flask import send_file
from docx import Document
from templates.reports import generate_report

# from app import app


from flask_login import login_required, login_user, logout_user, current_user
from models import User, db
from forms import SignUpForm, LoginForm


from werkzeug.security import check_password_hash, generate_password_hash


main = Blueprint('main', __name__)
# Project Routes

# @main.route('/')
# def index():
#     project_count = Project.query.count()
#     total_budget = db.session.query(db.func.sum(Project.project_amount)).scalar()
    
#     # More complex queries can be added here for detailed analysis
    
#     return render_template('landing_page.html', project_count=project_count, total_budget=total_budget)

@main.route('/')
def landing_page():
    try:
        # Simulate an operation that could fail, for illustration purposes
        # For example, fetching user-specific data from a database
        # If an error occurs, you could log it and/or take other appropriate actions
        return render_template('landing_page.html')
    except Exception as e:
        # Log the error here (logging not shown for brevity)
        print(f"An error occurred: {e}")  # Placeholder for actual logging

        # Optionally, you can return a custom error message or template
        # For now, we'll use Flask's abort function to return a 500 error
        # This will use your custom 500 error page if you have set up error handling for it
        abort(500)


@main.errorhandler(404)
def not_found_error(error):
    return render_template('404.html'), 404  # Assuming you have a 404.html template

@main.errorhandler(500)
def internal_error(error):
    # It's a good practice to roll back the session in case a database error caused the 500 error
    # db.session.rollback()
    return render_template('500.html'), 500  # Assuming you have a 500.html template



@main.route('/index')
def index():
    project_count = Project.query.count()

    # Fetch all projects to display in the projects list section
    projects = Project.query.all()

    # Example conversion rates; these should be dynamically fetched
    exchange_rates = {'EUR': 0.85, 'USD': 1, 'NGN': 900}

    # Example total budget in USD
    total_budget_usd = db.session.query(db.func.sum(Project.project_amount)).scalar() or 0

    # Convert total budget to other currencies
    total_budget_eur = total_budget_usd * exchange_rates['EUR']
    total_budget_ngn = total_budget_usd * exchange_rates['NGN']

    # Format the figures with commas for thousands and two decimal places
    formatted_total_budget_usd = "{:,.2f}".format(total_budget_usd)
    formatted_total_budget_eur = "{:,.2f}".format(total_budget_eur)
    formatted_total_budget_ngn = "{:,.2f}".format(total_budget_ngn)
    

    return render_template('dashboard.html', 
                           project_count=project_count,
                           projects=projects,
                           total_budget_usd=formatted_total_budget_usd,
                           total_budget_eur=formatted_total_budget_eur,
                           total_budget_ngn=formatted_total_budget_ngn)
    # return render_template('dashboard.html')


# @main.route('/', methods=['GET', 'POST'])
# def landing_page():
#     form = LoginForm()  # Instantiate the LoginForm

#     if form.validate_on_submit():
#         # Perform any necessary actions when the form is submitted
#         return redirect(url_for('main.dashboard'))

#     return render_template('landing_page.html', form=form)
    

@main.route('/activity_list', methods=['GET', 'POST'])
@login_required
def activity_list():

    activities = Activity.query.all()  # Fetch all activities

    return render_template('activity_list.html', activities=activities)

@main.route('/report_list', methods=['GET', 'POST'])
@login_required
def report_list():

    reports = Report.query.all()  # Fetch all activities

    return render_template('report_list.html', reports=reports)

# @main.route('/project', methods=['GET', 'POST'])
# @login_required
# def project():
#     if request.method == 'POST':
#         new_project = Project(
#             name=request.form['name'],
#             donor=request.form['donor'],
#             thematic_area=request.form['thematic_area'],
#             start_date=datetime.strptime(request.form['start_date'], '%Y-%m-%d'),
#             end_date=datetime.strptime(request.form['end_date'], '%Y-%m-%d'),
#             project_amount=float(request.form['project_amount'])
#         )
#         db.session.add(new_project)
#         db.session.commit()
#         return redirect(url_for('main.project'))

#     projects = Project.query.all()
#     return render_template('project.html', projects=projects)

@main.route('/create_project', methods=['GET', 'POST'])
@login_required
def create_project():
    if request.method == 'POST':
        try:
            new_project = Project(
                name=request.form['name'],
                donor=request.form['donor'],
                thematic_area=request.form['thematic_area'],
                start_date=datetime.strptime(request.form['start_date'], '%Y-%m-%d'),
                end_date=datetime.strptime(request.form['end_date'], '%Y-%m-%d'),
                project_amount=float(request.form['project_amount'])
            )
            db.session.add(new_project)
            db.session.commit()
            flash('Project added successfully!', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'An error occurred: {e}', 'danger')
            # Logging the error can be done here if you have a logging setup
        finally:
            return redirect(url_for('main.project'))

    projects = Project.query.all()
    return render_template('create_project.html', projects=projects)

@main.route('/project', methods=['GET', 'POST'])
@login_required
def project():
    if request.method == 'POST':
        try:
            new_project = Project(
                name=request.form['name'],
                donor=request.form['donor'],
                thematic_area=request.form['thematic_area'],
                start_date=datetime.strptime(request.form['start_date'], '%Y-%m-%d'),
                end_date=datetime.strptime(request.form['end_date'], '%Y-%m-%d'),
                project_amount=float(request.form['project_amount'])
            )
            db.session.add(new_project)
            db.session.commit()
            flash('Project added successfully!', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'An error occurred: {e}', 'danger')
            # Logging the error can be done here if you have a logging setup
        finally:
            return redirect(url_for('main.project'))

    projects = Project.query.all()
    return render_template('project.html', projects=projects)





# # Activities Routes
# @main.route('/activity', methods=['GET', 'POST'])
# @login_required
# def activity():
#     if request.method == 'POST':
#         project_id = request.form['project_id']  # Assuming project_id is passed in the form
#         new_activity = Activity(
#             project_id=project_id,
#             target_beneficiaries=request.form['target_beneficiaries'],
#             target_beneficiaries_male=int(request.form['target_beneficiaries_male']),
#             target_beneficiaries_female=int(request.form['target_beneficiaries_female']),
#             budget_amount=float(request.form['budget_amount']),
#             approved_budget_amount=float(request.form['approved_budget_amount'])
#         )
#         db.session.add(new_activity)
#         db.session.commit()
#         return redirect(url_for('main.activity'))

#     activities = Activity.query.all()
#     return render_template('activity.html', activities=activities)



# @main.route('/activity', methods=['GET', 'POST'])
# @login_required
# def activity():
#     if request.method == 'POST':
#         try:
#             # Validate or transform data as needed before creating a new activity
#             project_id = request.form.get('project_id')
#             name = request.form.get('name')
#             target_beneficiaries = request.form.get('target_beneficiaries')
#             target_beneficiaries_male = int(request.form.get('target_beneficiaries_male', 0))
#             target_beneficiaries_female = int(request.form.get('target_beneficiaries_female', 0))
#             budget_amount = float(request.form.get('budget_amount', 0))
#             approved_budget_amount = float(request.form.get('approved_budget_amount', 0))
            
#             # Assuming project_id is valid and exists; add validation as needed
#             new_activity = Activity(
#                 project_id=project_id,
#                 name=name,
#                 target_beneficiaries=target_beneficiaries,
#                 target_beneficiaries_male=target_beneficiaries_male,
#                 target_beneficiaries_female=target_beneficiaries_female,
#                 budget_amount=budget_amount,
#                 approved_budget_amount=approved_budget_amount
#             )
#             db.session.add(new_activity)
#             db.session.commit()
#             flash('New activity added successfully!', 'success')
#         except Exception as e:
#             # Log the exception here
#             db.session.rollback()  # Important to avoid session issues after an error
#             flash('Error adding activity. Please try again.', 'danger')
#             # Optionally log e or send it to an error tracking system
            
#         return redirect(url_for('main.activity'))

#     activities = Activity.query.all()
#     projects = Project.query.all()  # Fetch projects to populate the project_id select field
#     return render_template('activity.html', activities=activities, projects=projects)



@main.route('/activity', methods=['GET', 'POST'])
@login_required
def activity():
    form = ActivityForm()
    form.project_id.choices = [(p.id, p.name) for p in Project.query.all()]  # Populate project choices

    if form.validate_on_submit():
        try:
            new_activity = Activity(
                project_id=form.project_id.data,
                name=form.name.data,
                target_beneficiaries=form.target_beneficiaries.data,
                target_beneficiaries_male=form.target_beneficiaries_male.data,
                target_beneficiaries_female=form.target_beneficiaries_female.data,
                budget_amount=form.budget_amount.data,
                approved_budget_amount=form.approved_budget_amount.data
            )
            db.session.add(new_activity)
            db.session.commit()
            flash('New activity added successfully!', 'success')
            return redirect(url_for('main.activity'))
        except Exception as e:
            db.session.rollback()
            flash('Error adding activity. Please try again.', 'danger')

    activities = Activity.query.all()
    return render_template('activity.html', form=form, activities=activities)


@main.route('/report', methods=['GET', 'POST'])
@login_required
def report():
    if request.method == 'POST':
        activity_id = request.form['activity_id']
        new_report = Report(
            activity_id=activity_id,
            report_title=str(request.form['report_title']),
            date_from=datetime.strptime(request.form['date_from'], '%Y-%m-%d'),
            date_to=datetime.strptime(request.form['date_to'], '%Y-%m-%d'),
            number_reached_male=int(request.form['number_reached_male']),
            number_reached_female=int(request.form['number_reached_female']),
            written_report=request.files['written_report'].filename,
            photos=request.files['photos'].filename
        )
        db.session.add(new_report)
        db.session.commit()
        return redirect(url_for('main.report'))

    reports = Report.query.all()
    return render_template('report.html', reports=reports)







# Financial Report Routes
@main.route('/financial_report', methods=['GET', 'POST'])
@login_required
def financial_report():
    if request.method == 'POST':
        project_id = request.form['project_id']  # Assuming project_id is passed in the form
        new_financial_report = FinancialReport(
            project_id=project_id,
            budget=float(request.form['budget']),
            expenditure=float(request.form['expenditure'])
        )
        db.session.add(new_financial_report)
        db.session.commit()
        return redirect(url_for('main.financial_report'))

    financial_reports = FinancialReport.query.all()
    return render_template('financial_report.html', financial_reports=financial_reports)



@main.route('/analysis')
@login_required
def analysis():
    # Example of fetching data for analysis
    project_count = Project.query.count()
    total_budget = db.session.query(db.func.sum(Project.project_amount)).scalar()
    
    # More complex queries can be added here for detailed analysis
    
    return render_template('analysis.html',
                           project_count=project_count,
                           total_budget=total_budget,
                           total_budget_usd=total_budget_usd,
                           total_budget_eur=total_budget_eur,
                           total_budget_ngn=total_budget_ngn)


##############################
# Generate pdf report routes #
##############################

@main.route('/generate_pdf_project_report')
@login_required
def generate_pdf_project_report():
    projects = Project.query.all()
    html = render_template('pdf_report_project_template.html', projects=projects)
    
    return render_pdf(HTML(string=html))


@main.route('/generate_pdf_activity_report')
@login_required
def generate_pdf_activity_report():
    activities = Activity.query.all()
    project = Project.query.all()
    
    html = render_template('reports/pdf_activity_report_template.html', activities=activities, project=project) 

    return render_pdf(HTML(string=html))

@main.route('/generate_pdf_report')
@login_required
def generate_pdf_report():
    reports = Report.query.all()
    
    html = render_template('reports/pdf_report_template.html', reports=reports,) 

    return render_pdf(HTML(string=html))

# @main.route('/generate_pdf_report')
# @login_required
# def generate_pdf_report():
#     # Fetch data for the report
#     projects = Project.query.all()
#     activities = Activity.query.all()
#     reports = Report.query.all()
#     financial_reports = FinancialReport.query.all()
    
#     # Render a HTML template with this data
#     html = render_template('pdf_report_template.html', projects=projects, activities=activities, reports=reports, financial_reports=financial_reports)
    
#     # Convert HTML to PDF and return
#     return render_pdf(HTML(string=html))



# @main.route('/generate_pdf_report')
# @login_required
# def generate_pdf_report():
#     # Fetch data for the report
#     projects = Project.query.all()
    
#     # Render a HTML template with this data
#     html = render_template('pdf_report_template.html', projects=projects)
    
#     # Convert HTML to PDF
#     return render_pdf(HTML(string=html))


@main.route('/generate_excel_report')
@login_required
def generate_excel_report():
    # Fetch data for the report
    projects = Project.query.all()

    # Convert data to a DataFrame
    data = [{'donor': project.donor, 'thematic_area': project.thematic_area, 'start_date': project.start_date, 'end_date': project.end_date, 'project_amount': project.project_amount } for project in projects]  # Add all relevant fields
    df = pd.DataFrame(data)

    # Create an Excel file in memory
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Report')

    # Get the Excel file content
    excel_data = output.getvalue()

    # Return the Excel file in the response
    return Response(
        excel_data,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={"Content-Disposition": "attachment;filename=report.xlsx"}
    )


@main.route('/generate_word_report')
@login_required
def generate_word_report():
    # Fetch data for the report
    projects = Project.query.all()

    # Create a Word document
    doc = Document()
    doc.add_heading('Project Report', level=1)

    # Define table headers
    headers = ['Project ID', 'Project Name', 'Donor', 'Thematic Area', 'Project Start Date', 'Project End Date', 'Project Amount']

    # Add a table to the document
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'  # Add basic style to the table

    # Populate header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    # Add data to the table
    for project in projects:
        row_cells = table.add_row().cells
        row_cells[0].text = str(project.id)
        row_cells[1].text = project.name
        row_cells[2].text = project.donor  # Assuming 'donor' is a string
        row_cells[3].text = project.thematic_area  # Assuming 'thematic_area' is a string
        row_cells[4].text = project.start_date.strftime('%Y-%m-%d')  # Format the date
        row_cells[5].text = project.end_date.strftime('%Y-%m-%d')    # Format the date
        row_cells[6].text = str(project.project_amount)  # Convert float to string
    


    # Save the document in memory
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    # Return the Word document as a response
    return send_file(
        file_stream,
        as_attachment=True,
        download_name='project_report.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

# @main.route('/dashboard')
# @login_required
# def dashboard():
#     project_count = Project.query.count()
#     # Example conversion rates; these should be dynamically fetched
#     exchange_rates = {'EUR': 0.85, 'USD': 1, 'NGN': 900}

#     # Example total budget in USD
#     total_budget_usd = db.session.query(db.func.sum(Project.project_amount)).scalar()

#     # Initialize total_budget_eur and total_budget_ngn
#     total_budget_eur = 0
#     total_budget_ngn = 0

#     # Convert total budget to other currencies if total_budget_usd is not None
#     if total_budget_usd is not None:
#         total_budget_usd = total_budget_usd * exchange_rates['USD']
#         total_budget_eur = total_budget_usd * exchange_rates['EUR']
#         total_budget_ngn = total_budget_usd * exchange_rates['NGN']
#     else:
#         # Handle the case where total_budget_usd is None
#         print("Error: total_budget_usd is None")

#     return render_template('dashboard.html', 
#                            project_count=project_count,
#                            total_budget_usd=total_budget_usd,
#                            total_budget_eur=total_budget_eur,
#                            total_budget_ngn=total_budget_ngn)

# total_budget = db.session.query(db.func.sum(Project.project_amount)).scalar()


# @main.route('/dashboard')
# @login_required
# def dashboard():
#     # Fetch the count of projects
#     project_count = Project.query.count()

#     # Fetch all projects to display in the projects list section
#     projects = Project.query.all()

#     # Example conversion rates; these should be dynamically fetched
#     exchange_rates = {'EUR': 0.85, 'USD': 1, 'NGN': 900}

#     # Example total budget in USD
#     total_budget_usd = db.session.query(db.func.sum(Project.project_amount)).scalar()

#     # Initialize total_budget_eur and total_budget_ngn
#     total_budget_eur = 0
#     total_budget_ngn = 0

#     # Convert total budget to other currencies if total_budget_usd is not None
#     if total_budget_usd is not None:
#         total_budget_usd = total_budget_usd * exchange_rates['USD']  # This multiplication is redundant since the rate is 1
#         total_budget_eur = total_budget_usd * exchange_rates['EUR']
#         total_budget_ngn = total_budget_usd * exchange_rates['NGN']
#     else:
#         # Handle the case where total_budget_usd is None by setting default values
#         total_budget_usd = 0
#         print("Notice: No projects budget found, displaying defaults.")
        

#     return render_template('dashboard.html', 
#                            project_count=project_count,
#                            projects=projects,
#                            total_budget_usd=total_budget_usd,
#                            total_budget_eur=total_budget_eur,
#                            total_budget_ngn=total_budget_ngn)


@main.route('/dashboard')
@login_required
def dashboard():
    project_count = Project.query.count()

    # Fetch all projects to display in the projects list section
    projects = Project.query.all()

    # Example conversion rates; these should be dynamically fetched
    exchange_rates = {'EUR': 0.85, 'USD': 1, 'NGN': 900}

    # Example total budget in USD
    total_budget_usd = db.session.query(db.func.sum(Project.project_amount)).scalar() or 0

    # Convert total budget to other currencies
    total_budget_eur = total_budget_usd * exchange_rates['EUR']
    total_budget_ngn = total_budget_usd * exchange_rates['NGN']

    # Format the figures with commas for thousands and two decimal places
    formatted_total_budget_usd = "{:,.2f}".format(total_budget_usd)
    formatted_total_budget_eur = "{:,.2f}".format(total_budget_eur)
    formatted_total_budget_ngn = "{:,.2f}".format(total_budget_ngn)
    

    return render_template('dashboard.html', 
                           project_count=project_count,
                           projects=projects,
                           total_budget_usd=formatted_total_budget_usd,
                           total_budget_eur=formatted_total_budget_eur,
                           total_budget_ngn=formatted_total_budget_ngn)



@main.route('/signup', methods=['GET', 'POST'])
def signup():
    form = SignUpForm(request.form)

    if request.method == 'POST' and form.validate():
        # Check if user already exists
        existing_user = User.query.filter_by(email=form.email.data).first()
        if existing_user:
            flash('Email already registered.', 'danger')
            return redirect(url_for('main.signup'))

        # Hash password before storing
        hashed_password = generate_password_hash(form.password.data)

        new_user = User(
            name=form.name.data,
            email=form.email.data,
            phone=form.phone.data,
            password=hashed_password
        )
    try:
        db.session.add(new_user)
        db.session.commit()
        flash('Account created successfully. Please log in.', 'success')
        return redirect(url_for('main.login'))
    except Exception:
        db.session.rollback()
        flash('An error occurred. Please try again.', 'danger')

    return render_template('signup.html', form=form)



# @main.route('/login', methods=['GET', 'POST'])
# def login():
#     if current_user.is_authenticated:
#         flash("You are already logged in.", "info")
#         return redirect(url_for("main.index"))
#     form = LoginForm()
#     if form.validate_on_submit():
#         user = User.query.filter_by(email=form.email.data).first()
#         if user and check_password_hash(user.password, request.form["password"]):
#             login_user(user)
#             return redirect(url_for('main.login'))
        
#         # Check if user is admin
#         if user.is_admin:
#                 session['is_admin'] = True
#                 flash('Admin login successful', 'success')
#                 return redirect(url_for('/dashboard'))
#         else:
#             session['is_admin'] = False
#             flash('Login successful', 'success')

#         return redirect(url_for('/dashboard'))
    
#     else:
#         flash('Login Unsuccessful. Please check email and password', 'danger')

#     return render_template('/login.html', form=form)


@main.route('/login', methods=['GET', 'POST'])
def login():
    try:
        if current_user.is_authenticated:
            flash("You are already logged in.", "info")
            return redirect(url_for("main.index"))
        form = LoginForm()
        if form.validate_on_submit():
            user = User.query.filter_by(email=form.email.data).first()
            if user and check_password_hash(user.password, request.form["password"]):
                login_user(user)
                next_page = request.args.get('next')
                if user.is_admin:
                    session['is_admin'] = True
                    flash('Admin login successful', 'success')
                    return redirect(next_page or url_for('main.dashboard'))
                else:
                    session['is_admin'] = False
                    flash('Login successful', 'success')
                    return redirect(next_page or url_for('main.dashboard'))
            else:
                flash('Login Unsuccessful. Please check email and password', 'danger')
        return render_template('login.html', title='Login', form=form)
    except Exception as e:
        # Log the error here
        print(f"An error occurred during login: {e}")  # Use actual logging in production
        abort(500)  # Optionally, redirect to a custom error page


# Add a route for logging out
# @main.route('/logout')
# def logout():
#     session.pop('user_email', None)
#     flash('You have been logged out.', 'info')
#     return redirect(url_for('main.landing_page'))

@main.route('/logout')
# @login_required
def logout():
    logout_user()
    flash('You have been logged out successfully.', 'info')
    # return redirect(url_for('main.landing_page'))
    return render_template('landing_page.html')


from flask import jsonify

# Example route for fetching all projects
# @main.route('/api/projects')
# def get_projects():
#     projects = Project.query.all()
#     projects_data = [{'id': project.id, 'name': project.name} for project in projects]
#     return jsonify(projects_data)

@main.route('/api/projects')
def get_projects():
    projects = Project.query.all()
    projects_data = [{'id': project.id, 'name': project.name} for project in projects]
    return jsonify(projects_data)


# Example route for fetching activities by project ID
@main.route('/api/activities/<int:project_id>')
def get_activities(project_id):
    activities = Activity.query.filter_by(project_id=project_id).all()
    activities_data = [{'id': activity.id, 'name': activity.name} for activity in activities]
    return jsonify(activities_data)

# Example route for fetching budget amount by activity ID
@main.route('/api/activity/budget/<int:activity_id>')
def get_activity_budget(activity_id):
    activity = Activity.query.get_or_404(activity_id)
    return jsonify({'budget_amount': activity.budget_amount})


###################################
# Deleting and Editing routes     #
###################################


@main.route('/delete_activity/<int:activity_id>')
@login_required  
def delete_activity(activity_id):
    activity_to_delete = Activity.query.get(activity_id)
    if activity_to_delete:
        db.session.delete(activity_to_delete)
        db.session.commit()
        flash('Activity successfully removed', 'success')
    else:
        flash('Activity not found', 'error')
    return redirect(url_for('main.activity_list'))


@main.route('/edit_activity/<int:activity_id>', methods=['GET', 'POST'])
@login_required
def edit_activity(activity_id):
    activity = Activity.query.get_or_404(activity_id)
    form = ActivityForm(obj=activity)
    # Ensure this is done before form handling to repopulate on form re-render
    form.project_id.choices = [(p.id, p.name) for p in Project.query.all()]

    if form.validate_on_submit():
        # Form fields directly map to activity attributes
        activity.name = form.name.data
        activity.project_id = form.project_id.data  # Ensure this line is correct
        activity.target_beneficiaries = form.target_beneficiaries.data
        activity.target_beneficiaries_male = form.target_beneficiaries_male.data
        activity.target_beneficiaries_female = form.target_beneficiaries_female.data
        activity.budget_amount = form.budget_amount.data
        activity.approved_budget_amount = form.approved_budget_amount.data
        
        db.session.commit()
        flash('Activity updated successfully!', 'success')
        return redirect(url_for('main.activity_list'))

    return render_template('edit_activity.html', form=form, activity=activity)



# @main.route('/edit_activity/<int:activity_id>', methods=['GET', 'POST'])
# @login_required
# def edit_activity(activity_id):
#     activity = Activity.query.get_or_404(activity_id)
#     if request.method == 'POST':
#         activity.name = request.form.get('name')
#         activity.target_beneficiaries = request.form.get('target_beneficiaries')
#         activity.target_beneficiaries_male = int(request.form.get('target_beneficiaries_male', 0))
#         activity.target_beneficiaries_female = int(request.form.get('target_beneficiaries_female', 0))
#         activity.budget_amount = float(request.form.get('budget_amount', 0))
#         activity.approved_budget_amount = float(request.form.get('approved_budget_amount', 0))

#         db.session.commit()
#         flash('Activity updated successfully!', 'success')
#         return redirect(url_for('main.activity_list'))  # Ensure this redirects to your intended route
    
#     # For GET requests or if there's an issue with the POST request
#     # Assuming you have an 'edit_activity.html' template for the form
#     return render_template('edit_activity.html', activity=activity)


# @main.route('/edit_activity/<int:activity_id>', methods=['GET', 'POST'])
# @login_required
# def edit_activity(activity_id):
#     activity = Activity.query.get_or_404(activity_id)
#     form = ActivityForm(obj=activity)  # Pre-populate form with activity data

#     if form.validate_on_submit():
#         # Update the activity with form data
#         activity.name = form.name.data
#         activity.target_beneficiaries = form.target_beneficiaries.data
#         activity.target_beneficiaries_male = form.target_beneficiaries_male.data
#         activity.target_beneficiaries_female = form.target_beneficiaries_female.data
#         activity.budget_amount = form.budget_amount.data
#         activity.approved_budget_amount = form.approved_budget_amount.data
        
#         db.session.commit()
#         flash('Activity updated successfully!', 'success')
#         return redirect(url_for('main.activity_list'))

#     # Make sure to pass the form variable to the template
#     return render_template('edit_activity.html', form=form)



# @main.route('/edit_activity/<int:activity_id>', methods=['GET', 'POST'])
# @login_required
# def edit_activity(activity_id):
#     activity = Activity.query.get_or_404(activity_id)
#     form = ActivityForm(obj=activity)

#     form.project_id.choices = [(p.id, p.name) for p in Project.query.all()]  # Populate project choices


#     if form.validate_on_submit():
#         # Update the activity with form data
#         activity.name = form.name.data
#         activity.target_beneficiaries = form.target_beneficiaries.data
#         activity.target_beneficiaries_male = form.target_beneficiaries_male.data
#         activity.target_beneficiaries_female = form.target_beneficiaries_female.data
#         activity.budget_amount = form.budget_amount.data
#         activity.approved_budget_amount = form.approved_budget_amount.data
        
        
#         db.session.commit()
#         flash('Activity updated successfully!', 'success')
#         return redirect(url_for('main.activity_list'))

#     return render_template('edit_activity.html', form=form, activity=activity)



# @main.route('/edit_activity/<int:activity_id>', methods=['GET', 'POST'])
# @login_required
# def edit_activity(activity_id):
#     activity = Activity.query.get_or_404(activity_id)
#     if request.method == 'POST':
#         activity.name = request.form.name.data
#         activity.target_beneficiaries = request.form.target_beneficiaries.data
#         activity.target_beneficiaries_male = request.form.target_beneficiaries_male.data
#         activity.target_beneficiaries_female = request.form.target_beneficiaries_female.data
#         activity.budget_amount = request.form.budget_amount.data
#         activity.approved_budget_amount = request.form.approved_budget_amount.data

#         db.session.commit()
#         flash('Activity updated successfully!', 'success')
#         return redirect(url_for('main.activity_list'))  # Adjust the redirect as needed
    
#     return render_template('activity_list.html', activity=activity)


    # return render_template('edit_project.html', project=project)

    # form = ActivityForm(obj=activity)
    
    # if form.validate_on_submit():
    #     activity.name = form.name.data
    #     activity.target_beneficiaries = form.target_beneficiaries.data
    #     activity.target_beneficiaries_male = form.target_beneficiaries_male.data
    #     activity.target_beneficiaries_female = form.target_beneficiaries_female.data
    #     activity.budget_amount = form.budget_amount.data
    #     activity.approved_budget_amount = form.approved_budget_amount.data
    #     db.session.commit()
    #     flash('Activity updated successfully!', 'success')
    #     return redirect(url_for('main.activity_list'))  # Adjust the redirect as needed
    
    # return render_template('activity_list.html', form=form)
    

@main.route('/edit_report/<int:report_id>', methods=['GET', 'POST'])
def edit_report(report_id):
    report = Report.query.get_or_404(report_id)

    written_report_path = None
    photos_path = None

    if request.method == 'POST':
        # Assuming your report model has fields such as donor, date_from, etc.
        report.donor = request.form['donor']
        report.start_date = datetime.strptime(request.form['date_from'], '%Y-%m-%d')
        report.end_date = datetime.strptime(request.form['date_to'], '%Y-%m-%d')
        report.number_reached_male=int(request.form['number_reached_male'])
        report.number_reached_female=int(request.form['number_reached_female'])
        report.written_report=written_report_path
        report.photos=photos_path
        # ... handle other fields similarly

        db.session.commit()
        flash('Report updated successfully!', 'success')
        return redirect(url_for('main.report'))  # Replace with your actual list view route

    return render_template('edit_report.html', report=report)



@main.route('/delete_report/<int:report_id>', methods=['POST'])
def delete_report(report_id):
    report = Report.query.get_or_404(report_id)

    db.session.delete(report)
    db.session.commit()
    flash('Report deleted successfully!', 'info')
    return redirect(url_for('main.report'))  # Replace with your actual list view route


@main.route('/edit_project/<int:project_id>', methods=['GET', 'POST'])
def edit_project(project_id):
    project = Project.query.get_or_404(project_id)

    if request.method == 'POST':
        project.name = request.form['name']
        project.donor = request.form['donor']
        project.thematic_area = request.form['thematic_area']
        project.start_date=datetime.strptime(request.form['start_date'], '%Y-%m-%d')
        project.end_date = datetime.strptime(request.form['end_date'], '%Y-%m-%d')
        project.project_amount = float(request.form['project_amount'])

        db.session.commit()
        return redirect(url_for('main.project'))  # Replace with your view function

    return render_template('edit_project.html', project=project)



@main.route('/delete_project/<int:project_id>', methods=['POST'])
def delete_project(project_id):
    project = Project.query.get_or_404(project_id)
    db.session.delete(project)
    db.session.commit()
    return redirect(url_for('main.project'))