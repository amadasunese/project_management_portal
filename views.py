from flask import Blueprint, render_template, request, redirect, url_for, abort, flash, session
from flask_weasyprint import HTML, render_pdf
from models import db
from models import Project, Activity, Report, FinancialReport
from forms import project_choices, ProjectForm, ReportForm, ActivityForm, FinancialReportForm
from datetime import datetime
from werkzeug.utils import secure_filename
import os
from flask import Response
import pandas as pd
from io import BytesIO
from flask import send_file
from docx import Document
from templates.reports import generate_report
from werkzeug.security import check_password_hash, generate_password_hash
from flask_login import login_required, login_user, logout_user, current_user
from models import User, db
from forms import SignUpForm, LoginForm
from flask import jsonify

import matplotlib.pyplot as plt

from flask import Flask, Response
from matplotlib.backends.backend_agg import FigureCanvasAgg as FigureCanvas
from matplotlib.figure import Figure
import io



main = Blueprint('main', __name__)

@main.route('/')
def landing_page():
    try:

        return render_template('landing_page.html')
    except Exception as e:
        print(f"An error occurred: {e}")
        abort(500)


@main.errorhandler(404)
def not_found_error(error):
    return render_template('404.html'), 404

@main.errorhandler(500)
def internal_error(error):
    return render_template('500.html'), 500


@main.route('/index')
def index():
    project_count = Project.query.count()

    projects = Project.query.all()

    exchange_rates = {'EUR': 0.85, 'USD': 1, 'NGN': 900}

    total_budget_usd = db.session.query(db.func.sum(Project.project_amount)).scalar() or 0

    total_budget_eur = total_budget_usd * exchange_rates['EUR']
    total_budget_ngn = total_budget_usd * exchange_rates['NGN']

    formatted_total_budget_usd = "{:,.2f}".format(total_budget_usd)
    formatted_total_budget_eur = "{:,.2f}".format(total_budget_eur)
    formatted_total_budget_ngn = "{:,.2f}".format(total_budget_ngn)


    return render_template('dashboard.html',
                           project_count=project_count,
                           projects=projects,
                           total_budget_usd=formatted_total_budget_usd,
                           total_budget_eur=formatted_total_budget_eur,
                           total_budget_ngn=formatted_total_budget_ngn)


@main.route('/activity_list', methods=['GET', 'POST'])
@login_required
def activity_list():
    activities = Activity.query.all()
    return render_template('activity_list.html', activities=activities)

@main.route('/report_list', methods=['GET', 'POST'])
@login_required
def report_list():
    reports = Report.query.all()
    return render_template('report_list.html', reports=reports)


@main.route('/create_project', methods=['GET', 'POST'])
@login_required
def create_project():
    if request.method == 'POST':
        try:
            new_project = Project(
                name=request.form['name'],
                donor=request.form['donor'],
                thematic_area=request.form['thematic_area'],
                start_date=datetime.strptime(request.form['start_date'], '%Y-%m-%d'),
                end_date=datetime.strptime(request.form['end_date'], '%Y-%m-%d'),
                project_amount=float(request.form['project_amount'])
            )
            db.session.add(new_project)
            db.session.commit()
            flash('Project added successfully!', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'An error occurred: {e}', 'danger')
        finally:
            return redirect(url_for('main.project'))

    projects = Project.query.all()
    return render_template('create_project.html', projects=projects)

@main.route('/project', methods=['GET', 'POST'])
@login_required
def project():
    if request.method == 'POST':
        try:
            new_project = Project(
                name=request.form['name'],
                donor=request.form['donor'],
                thematic_area=request.form['thematic_area'],
                start_date=datetime.strptime(request.form['start_date'], '%Y-%m-%d'),
                end_date=datetime.strptime(request.form['end_date'], '%Y-%m-%d'),
                project_amount=float(request.form['project_amount'])
            )
            db.session.add(new_project)
            db.session.commit()
            flash('Project added successfully!', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'An error occurred: {e}', 'danger')
        finally:
            return redirect(url_for('main.project'))

    projects = Project.query.all()
    return render_template('project.html', projects=projects)

@main.route('/project_list', methods=['GET', 'POST'])
@login_required
def project_list():
    projects = Project.query.all()
    return render_template('project_list.html', projects=projects)

@main.route('/activity', methods=['GET', 'POST'])
@login_required
def activity():
    form = ActivityForm()
    form.project_id.choices = [(p.id, p.name) for p in Project.query.all()]

    if form.validate_on_submit():
        try:
            new_activity = Activity(
                project_id=form.project_id.data,
                name=form.name.data,
                target_beneficiaries=form.target_beneficiaries.data,
                target_beneficiaries_male=form.target_beneficiaries_male.data,
                target_beneficiaries_female=form.target_beneficiaries_female.data,
                budget_amount=form.budget_amount.data,
                approved_budget_amount=form.approved_budget_amount.data
            )
            db.session.add(new_activity)
            db.session.commit()
            flash('New activity added successfully!', 'success')
            return redirect(url_for('main.activity'))
        except Exception as e:
            db.session.rollback()
            flash('Error adding activity. Please try again.', 'danger')

    activities = Activity.query.all()
    return render_template('activity.html', form=form, activities=activities)

UPLOAD_FOLDER = 'static/uploads'
ALLOWED_EXTENSIONS = {'txt', 'doc', 'docx', 'pdf', 'png', 'jpg', 'jpeg', 'gif'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@main.route('/report', methods=['GET', 'POST'])
@login_required
def report():
    if request.method == 'POST':
        if 'written_report' in request.files and 'photos' in request.files:
            written_report_file = request.files['written_report']
            photos_file = request.files['photos']

            if written_report_file.filename == '' or photos_file.filename == '':
                flash('No selected file')
                return redirect(request.url)

            if written_report_file and allowed_file(written_report_file.filename) and photos_file and allowed_file(photos_file.filename):
                written_report_filename = secure_filename(written_report_file.filename)
                photos_filename = secure_filename(photos_file.filename)

                written_report_path = os.path.join(UPLOAD_FOLDER, written_report_filename)
                photos_path = os.path.join(UPLOAD_FOLDER, photos_filename)

                written_report_file.save(written_report_path)
                photos_file.save(photos_path)

                new_report = Report(
                    activity_id=request.form['activity_id'],
                    report_title=request.form['report_title'],
                    date_from=datetime.strptime(request.form['date_from'], '%Y-%m-%d'),
                    date_to=datetime.strptime(request.form['date_to'], '%Y-%m-%d'),
                    number_reached_male=int(request.form['number_reached_male']),
                    number_reached_female=int(request.form['number_reached_female']),
                    written_report=written_report_path,
                    photos=photos_path
                )
                db.session.add(new_report)
                db.session.commit()
                flash('Report successfully added!')
            else:
                flash('Invalid file type.')
        else:
            flash('Missing file uploads.')

        return redirect(url_for('main.report'))

    reports = Report.query.all()
    return render_template('report.html', reports=reports)

@main.route('/financial_report', methods=['GET', 'POST'])
def financial_report():
    if request.method == 'POST':
        try:
            project_id = request.form['project_id']
            total_expenditure = float(request.form['total_expenditure'])
            financial_report = FinancialReport.query.filter_by(project_id=project_id).first()

            if financial_report:
                financial_report.total_expenditure = total_expenditure
            else:
                financial_report = FinancialReport(project_id=project_id, total_expenditure=total_expenditure)
                db.session.add(financial_report)

            db.session.commit()
            flash('Financial report submitted successfully.', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Error submitting financial report: {e}', 'error')
            return redirect(url_for('financial_report'))

    projects = Project.query.all()
    return render_template('financial_report.html', projects=projects)

@main.route('/get-budgets/<int:project_id>')
def get_budgets(project_id):
    project = Project.query.gety(project_id)
    if project:
        return jsonify({'budget': project.budget})
    return jsonify({'error': 'Project not found'}), 404


@main.route('/financial_reports_list', methods=['GET', 'POST'])
@login_required
def financial_reports_list():
    financial_reports = FinancialReport.query.all()
    return render_template('financial_reports_list.html', financial_reports=financial_reports)


@main.route('/fetch-budget', methods=['GET'])
def fetch_budget():
    project_id = request.args.get('projectId', type=int)
    if project_id is not None:
        project = Project.query.filter_by(id=project_id).first()
        if project:
            return jsonify({'budget': project.project_amount})
        else:
            return jsonify({'error': 'Project not found'}), 404
    return jsonify({'error': 'Invalid request'}), 400

@main.route('/get-projects', methods=['GET'])
@login_required
def get_project():
    projects = Project.query.all()
    projects_data = [{'id': project.id, 'name': project.name} for project in projects]
    return jsonify(projects_data)


@main.route('/get_project_budget/<int:project_id>')
def get_project_budget(project_id):
    project = Project.query.all()
    return jsonify(budget=project.project_amount)

# Search API
@main.route('/api/projects')
def get_projects():
    projects = Project.query.all()
    projects_data = [{'id': project.id, 'name': project.name} for project in projects]
    return jsonify(projects_data)

@main.route('/get-budget/<int:project_id>')
def get_budget(project_id):
    project = FinancialReport.query.filter_by(project_id=project_id).first()
    if project:
        return jsonify({'budget': project.budget})
    return jsonify({'error': 'Project not found'}), 404

def project_choices():
    return Project.query.all()


@main.route('/analysis')
@login_required
def analysis():
    project_count = Project.query.count()
    total_budget = db.session.query(db.func.sum(Project.project_amount)).scalar()

    total_budget_usd = db.session.query(db.func.sum(Project.project_amount)).scalar() or 0

    exchange_rates = {'EUR': 0.85, 'USD': 1, 'NGN': 900}

    total_budget_eur = total_budget_usd * exchange_rates['EUR']
    total_budget_ngn = total_budget_usd * exchange_rates['NGN']

    formatted_total_budget_usd = "{:,.2f}".format(total_budget_usd)
    formatted_total_budget_eur = "{:,.2f}".format(total_budget_eur)
    formatted_total_budget_ngn = "{:,.2f}".format(total_budget_ngn)

    return render_template('analysis.html',
                           project_count=project_count,
                           total_budget=total_budget,
                           total_budget_usd=total_budget_usd,
                           total_budget_eur=total_budget_eur,
                           total_budget_ngn=total_budget_ngn)


##############################
# Generate pdf report routes #
##############################

@main.route('/generate_pdf_project_report')
@login_required
def generate_pdf_project_report():
    projects = Project.query.all()
    html = render_template('pdf_report_project_template.html', projects=projects)

    return render_pdf(HTML(string=html))


@main.route('/generate_pdf_activity_report')
@login_required
def generate_pdf_activity_report():
    activities = Activity.query.all()
    project = Project.query.all()

    html = render_template('reports/pdf_activity_report_template.html', activities=activities, project=project)

    return render_pdf(HTML(string=html))

@main.route('/generate_pdf_report')
@login_required
def generate_pdf_report():
    reports = Report.query.all()

    html = render_template('reports/pdf_report_template.html', reports=reports,)

    return render_pdf(HTML(string=html))

@main.route('/generate_excel_report')
@login_required
def generate_excel_report():
    projects = Project.query.all()

    data = [{'donor': project.donor, 'thematic_area': project.thematic_area, 'start_date': project.start_date, 'end_date': project.end_date, 'project_amount': project.project_amount } for project in projects]  # Add all relevant fields
    df = pd.DataFrame(data)

    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='Report')
    excel_data = output.getvalue()
    return Response(
        excel_data,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={"Content-Disposition": "attachment;filename=report.xlsx"}
    )


@main.route('/generate_word_report')
@login_required
def generate_word_report():
    projects = Project.query.all()

    doc = Document()
    doc.add_heading('Project Report', level=1)

    headers = ['Project ID', 'Project Name', 'Donor', 'Thematic Area', 'Project Start Date', 'Project End Date', 'Project Amount']

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    for project in projects:
        row_cells = table.add_row().cells
        row_cells[0].text = str(project.id)
        row_cells[1].text = project.name
        row_cells[2].text = project.donor
        row_cells[3].text = project.thematic_area
        row_cells[4].text = project.start_date.strftime('%Y-%m-%d')
        row_cells[5].text = project.end_date.strftime('%Y-%m-%d')
        row_cells[6].text = str(project.project_amount)


    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name='project_report.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@main.route('/dashboard')
@login_required
def dashboard():
    project_count = Project.query.count()

    projects = Project.query.all()

    exchange_rates = {'EUR': 0.85, 'USD': 1, 'NGN': 900}

    total_budget_usd = db.session.query(db.func.sum(Project.project_amount)).scalar() or 0

    total_budget_eur = total_budget_usd * exchange_rates['EUR']
    total_budget_ngn = total_budget_usd * exchange_rates['NGN']

    formatted_total_budget_usd = "{:,.2f}".format(total_budget_usd)
    formatted_total_budget_eur = "{:,.2f}".format(total_budget_eur)
    formatted_total_budget_ngn = "{:,.2f}".format(total_budget_ngn)

    return render_template('dashboard.html',
                           project_count=project_count,
                           projects=projects,
                           total_budget_usd=formatted_total_budget_usd,
                           total_budget_eur=formatted_total_budget_eur,
                           total_budget_ngn=formatted_total_budget_ngn)


##############################
# Signup and login routes #
##############################

@main.route('/signup', methods=['GET', 'POST'])
def signup():
    form = SignUpForm(request.form)

    if request.method == 'POST' and form.validate():
        """Check if user already exists"""
        existing_user = User.query.filter_by(email=form.email.data).first()
        if existing_user:
            flash('Email already registered.', 'danger')
            return redirect(url_for('main.signup'))

        hashed_password = generate_password_hash(form.password.data)

        new_user = User(
            name=form.name.data,
            email=form.email.data,
            phone_number=form.phone_number.data,
            password=hashed_password
        )
    try:
        db.session.add(new_user)
        db.session.commit()
        flash('Account created successfully. Please log in.', 'success')
        return redirect(url_for('main.login'))
    except Exception:
        db.session.rollback()
        flash('An error occurred. Please try again.', 'danger')

    return render_template('signup.html', form=form)


@main.route('/login', methods=['GET', 'POST'])
def login():
    try:
        if current_user.is_authenticated:
            flash("You are already logged in.", "info")
            return redirect(url_for("main.index"))
        form = LoginForm()
        if form.validate_on_submit():
            user = User.query.filter_by(email=form.email.data).first()
            if user and check_password_hash(user.password, request.form["password"]):
                login_user(user)
                next_page = request.args.get('next')
                if user.is_admin:
                    session['is_admin'] = True
                    flash('Admin login successful', 'success')
                    return redirect(next_page or url_for('main.dashboard'))
                else:
                    session['is_admin'] = False
                    flash('Login successful', 'success')
                    return redirect(next_page or url_for('main.dashboard'))
            else:
                flash('Login Unsuccessful. Please check email and password', 'danger')
        return render_template('login.html', title='Login', form=form)
    except Exception as e:
        """Log the error here"""
        print(f"An error occurred during login: {e}")
        abort(500)


@main.route('/logout')
# @login_required
def logout():
    logout_user()
    flash('You have been logged out successfully.', 'info')
    return render_template('landing_page.html')


@main.route('/api/activities/<int:project_id>')
def get_activities(project_id):
    activities = Activity.query.filter_by(project_id=project_id).all()
    activities_data = [{'id': activity.id, 'name': activity.name} for activity in activities]
    return jsonify(activities_data)

@main.route('/api/activity/budget/<int:activity_id>')
def get_activity_budget(activity_id):
    activity = Activity.query.get_or_404(activity_id)
    return jsonify({'budget_amount': activity.budget_amount})


###################################
# Deleting and Editing routes     #
###################################


@main.route('/delete_activity/<int:activity_id>')
@login_required
def delete_activity(activity_id):
    activity_to_delete = Activity.query.get(activity_id)
    if activity_to_delete:
        db.session.delete(activity_to_delete)
        db.session.commit()
        flash('Activity successfully removed', 'success')
    else:
        flash('Activity not found', 'error')
    return redirect(url_for('main.activity_list'))


@main.route('/edit_activity/<int:activity_id>', methods=['GET', 'POST'])
@login_required
def edit_activity(activity_id):
    activity = Activity.query.get_or_404(activity_id)
    form = ActivityForm(obj=activity)

    form.project_id.choices = [(p.id, p.name) for p in Project.query.all()]

    if form.validate_on_submit():
        activity.name = form.name.data
        activity.project_id = form.project_id.data
        activity.target_beneficiaries = form.target_beneficiaries.data
        activity.target_beneficiaries_male = form.target_beneficiaries_male.data
        activity.target_beneficiaries_female = form.target_beneficiaries_female.data
        activity.budget_amount = form.budget_amount.data
        activity.approved_budget_amount = form.approved_budget_amount.data

        db.session.commit()
        flash('Activity updated successfully!', 'success')
        return redirect(url_for('main.activity_list'))

    return render_template('edit_activity.html', form=form, activity=activity)


@main.route('/edit_report/<int:report_id>', methods=['GET', 'POST'])
def edit_report(report_id):
    report = Report.query.get_or_404(report_id)

    written_report_path = None
    photos_path = None

    if request.method == 'POST':
        report.donor = request.form['donor']
        report.start_date = datetime.strptime(request.form['date_from'], '%Y-%m-%d')
        report.end_date = datetime.strptime(request.form['date_to'], '%Y-%m-%d')
        report.number_reached_male=int(request.form['number_reached_male'])
        report.number_reached_female=int(request.form['number_reached_female'])
        report.written_report=written_report_path
        report.photos=photos_path


        db.session.commit()
        flash('Report updated successfully!', 'success')
        return redirect(url_for('main.report'))

    return render_template('edit_report.html', report=report)



@main.route('/delete_report/<int:report_id>', methods=['GET', 'POST'])
def delete_report(report_id):
    report = Report.query.get_or_404(report_id)

    db.session.delete(report)
    db.session.commit()
    flash('Report deleted successfully!', 'info')
    return redirect(url_for('main.report'))


@main.route('/edit_project/<int:project_id>', methods=['GET', 'POST'])
def edit_project(project_id):
    project = Project.query.get_or_404(project_id)

    if request.method == 'POST':
        project.name = request.form['name']
        project.donor = request.form['donor']
        project.thematic_area = request.form['thematic_area']
        project.start_date=datetime.strptime(request.form['start_date'], '%Y-%m-%d')
        project.end_date = datetime.strptime(request.form['end_date'], '%Y-%m-%d')
        project.project_amount = float(request.form['project_amount'])

        db.session.commit()
        return redirect(url_for('main.project'))

    return render_template('edit_project.html', project=project)



@main.route('/delete_project/<int:project_id>', methods=['POST'])
def delete_project(project_id):
    project = Project.query.get_or_404(project_id)
    db.session.delete(project)
    db.session.commit()
    return redirect(url_for('main.project'))


def get_project_data():
    from models import Project, Activity
    project_names = []
    expenditure_percentages = []

    projects = Project.query.all()
    for project in projects:
        if project.project_amount > 0:
            total_approved_budget = sum(activity.approved_budget_amount for activity in project.activities)
            expenditure_percentage = (total_approved_budget / project.project_amount) * 100

            project_names.append(project.name)
            expenditure_percentages.append(expenditure_percentage)

    return project_names, expenditure_percentages


@main.route('/dynamic_chart.png')
def dynamic_chart():
    project_names, expenditure_percentages = get_project_data()

    fig = Figure()
    axis = fig.add_subplot(1, 1, 1)
    axis.bar(project_names, expenditure_percentages)
    axis.set_title('Project Expenditure as a Percentage of Total Budget')
    axis.set_ylabel('Percentage')
    axis.set_xticks(range(len(project_names)))
    axis.set_xticklabels(project_names, rotation=45, ha="right")

    # Convert plot to PNG image
    png_image = io.BytesIO()
    FigureCanvas(fig).print_png(png_image)

    # Return PNG image as response
    return Response(png_image.getvalue(), mimetype='image/png')